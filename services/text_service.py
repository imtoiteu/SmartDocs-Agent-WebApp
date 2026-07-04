"""Read plain text from TXT / DOCX / PDF uploads."""

def read_txt(path):
    with open(path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()

def read_docx(path):
    from docx import Document
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

def read_pdf_text(path):
    import pypdfium2 as pdfium
    doc = pdfium.PdfDocument(path)
    pages = []
    for page in doc:
        tp = page.get_textpage()
        pages.append(tp.get_text_range())
        tp.close(); page.close()
    doc.close()
    return "\n\n".join(pages)

def read_file(path: str) -> str:
    p = path.lower()
    if p.endswith(".txt"):  return read_txt(path)
    if p.endswith(".docx"): return read_docx(path)
    if p.endswith(".pdf"):  return read_pdf_text(path)
    raise ValueError(f"Unsupported file type: {path}")
